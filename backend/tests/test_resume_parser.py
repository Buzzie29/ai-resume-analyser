from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from app.services.resume_parser import extract_resume_text

TESTS_DIR = Path(__file__).parent
TEST_RESUME = TESTS_DIR / "sample_resume.txt"


def test_extract_resume_text_from_txt():
    text = extract_resume_text(str(TEST_RESUME))

    assert text
    assert "Tanmay" in text
    assert "Python" in text


def test_extract_resume_text_from_docx(tmp_path):
    file_path = tmp_path / "resume.docx"

    document = Document()
    document.add_paragraph("Tanmay Kohale")
    document.add_paragraph("Python Developer")
    document.add_paragraph("Skills: Python, FastAPI, Git")

    document.save(file_path)

    text = extract_resume_text(str(file_path))

    assert "Tanmay Kohale" in text
    assert "Python Developer" in text
    assert "FastAPI" in text


def test_extract_resume_text_from_pdf(tmp_path):
    file_path = tmp_path / "resume.pdf"

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(file_path)

    text = extract_resume_text(str(file_path))

    assert isinstance(text, str)
